#!/usr/bin/env python3
"""
Генератор .docx с дефолтами оформления по ГОСТ Р 7.0.97-2025.

Зашитые параметры:
- формат A4, поля: левое 20 / правое 10 / верхнее 20 / нижнее 20 мм
  (long_term_storage=True -> левое 30 мм для хранения свыше 10 лет);
- шрифт Times New Roman 14 (допустимо 12/13/14);
- абзацный отступ 1,25 см; межстрочный интервал 1,5; выравнивание по ширине;
- нумерация страниц со 2-й, посередине верхнего поля.

Зависимость: python-docx  (pip install python-docx)

Использование как библиотеки:

    from gost_docx import GostDocument
    d = GostDocument(font_size=14)
    d.org_name("АКЦИОНЕРНОЕ ОБЩЕСТВО \"ПРИМЕР\"", short="АО \"Пример\"")
    d.doc_kind("ПРИКАЗ")            # реквизит 09 (кроме писем)
    d.date_and_number("05.06.2024", "12")
    d.place("г. Москва")           # реквизит 13
    d.title("О создании комиссии") # реквизит 17
    d.body_paragraph("В целях ...  приказываю:")
    d.numbered("Создать комиссию в составе ...")
    d.signature("Генеральный директор", "И. И. Иванов")
    d.executor("Петров Пётр Петрович", "+7 495 000-00-00")
    d.save("prikaz.docx")

Скрипт задаёт форму по ГОСТу. Содержание текста и его «человечность» —
на стороне вызывающего (см. references/ai-tells-ru.md).
"""

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class GostDocument:
    def __init__(self, font_name="Times New Roman", font_size=14,
                 long_term_storage=False, line_spacing=1.5):
        if font_size not in (12, 13, 14):
            raise ValueError("Размер шрифта по ГОСТу: 12, 13 или 14")
        self.font_name = font_name
        self.font_size = font_size
        self.doc = Document()
        self._setup_page(long_term_storage)
        self._setup_normal_style(line_spacing)
        self._add_page_numbers()

    # --- базовая настройка страницы и стиля ---

    def _setup_page(self, long_term_storage):
        left = Mm(30) if long_term_storage else Mm(20)
        for section in self.doc.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.left_margin = left
            section.right_margin = Mm(10)
            section.top_margin = Mm(20)
            section.bottom_margin = Mm(20)

    def _setup_normal_style(self, line_spacing):
        style = self.doc.styles["Normal"]
        style.font.name = self.font_name
        # кириллица требует явного указания гарнитуры для восточноазиатских/cs
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), self.font_name)
        rfonts.set(qn("w:hAnsi"), self.font_name)
        rfonts.set(qn("w:cs"), self.font_name)
        style.font.size = Pt(self.font_size)
        pf = style.paragraph_format
        pf.first_line_indent = Cm(1.25)
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(0)

    def _add_page_numbers(self):
        # номер страницы посередине верхнего поля; на 1-й странице скрыт
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        fldchar1 = OxmlElement("w:fldChar")
        fldchar1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fldchar2 = OxmlElement("w:fldChar")
        fldchar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldchar1)
        run._r.append(instr)
        run._r.append(fldchar2)

    # --- вспомогательное ---

    def _para(self, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None,
              bold=False, space_after=0, caps=False):
        p = self.doc.add_paragraph()
        p.alignment = align
        if indent is not None:
            p.paragraph_format.first_line_indent = indent
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text.upper() if caps else text)
            run.bold = bold
        return p

    # --- реквизиты (наиболее частые) ---

    def org_name(self, full, short=None, parent=None):
        """05 — наименование организации (+ вышестоящая, + сокращённое)."""
        if parent:
            self._para(parent, align=WD_ALIGN_PARAGRAPH.CENTER, indent=Cm(0))
        self._para(full, align=WD_ALIGN_PARAGRAPH.CENTER, indent=Cm(0), bold=True)
        if short:
            self._para("(%s)" % short, align=WD_ALIGN_PARAGRAPH.CENTER,
                       indent=Cm(0), space_after=6)

    def doc_kind(self, kind):
        """09 — наименование вида документа (кроме писем), прописными."""
        self._para(kind, align=WD_ALIGN_PARAGRAPH.CENTER, indent=Cm(0),
                   bold=True, caps=True, space_after=6)

    def date_and_number(self, date, number, place_left=True):
        """10 и 11 — дата и регистрационный номер в одной строке."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        tab_stops = p.paragraph_format.tab_stops
        p.add_run("%s   № %s" % (date, number))
        return p

    def place(self, place):
        """13 — место составления."""
        self._para(place, align=WD_ALIGN_PARAGRAPH.CENTER, indent=Cm(0),
                   space_after=12)

    def addressee(self, lines):
        """15 — адресат: справа сверху, флаговое выравнивание."""
        for i, line in enumerate(lines):
            self._para(line, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(9),
                       space_after=0 if i < len(lines) - 1 else 12)

    def approval(self, lines):
        """16 — гриф утверждения: правый верхний угол."""
        for i, line in enumerate(lines):
            self._para(line, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(9),
                       space_after=0 if i < len(lines) - 1 else 12)

    def title(self, text):
        """17 — заголовок к тексту («О чём»), слева от левого поля."""
        self._para(text, align=WD_ALIGN_PARAGRAPH.LEFT, indent=Cm(0),
                   space_after=12)

    def body_paragraph(self, text, bold=False):
        """18 — абзац текста (абзацный отступ 1,25, по ширине)."""
        self._para(text, bold=bold)

    def numbered(self, text, number=None):
        """Пункт нумерованного текста (рубрикация арабскими цифрами)."""
        prefix = ("%s. " % number) if number else ""
        self._para(prefix + text)

    def signature(self, position, name):
        """22 — подпись: должность слева, И. О. Фамилия справа."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16),
                                                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        p.add_run("%s\t%s" % (position, name))
        return p

    def executor(self, fio, phone, unit=None, email=None):
        """25 — отметка об исполнителе: внизу слева, меньшим кеглем."""
        parts = [fio, phone]
        if unit:
            parts.insert(0, unit)
        if email:
            parts.append(email)
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(", ".join(parts))
        run.font.size = Pt(max(self.font_size - 2, 8))
        return p

    def raw(self):
        """Доступ к объекту python-docx Document для тонкой настройки."""
        return self.doc

    def save(self, path):
        self.doc.save(path)
        return path


if __name__ == "__main__":
    # демонстрационный приказ
    d = GostDocument(font_size=14)
    d.org_name('АКЦИОНЕРНОЕ ОБЩЕСТВО "ПРИМЕР"', short='АО "Пример"')
    d.doc_kind("ПРИКАЗ")
    d.date_and_number("05.06.2024", "12")
    d.place("г. Москва")
    d.title("О создании аттестационной комиссии")
    d.body_paragraph("Для проведения аттестации работников приказываю:")
    d.numbered("Создать аттестационную комиссию в составе [__].", 1)
    d.numbered("Контроль за исполнением приказа оставляю за собой.", 2)
    d.signature("Генеральный директор", "И. И. Иванов")
    d.executor("Петров Пётр Петрович", "+7 495 000-00-00",
               unit="Отдел кадров")
    out = d.save("prikaz_demo.docx")
    print("Сохранено:", out)
